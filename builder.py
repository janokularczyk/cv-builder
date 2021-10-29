from docx import Document
from docx.shared import Inches
import pyttsx3

def speak(text):
    pyttsx3.speak(text)

document = Document()

# profile picture
document.add_picture(
    'images/me.jpg',
    width=Inches(2.0)
)

# name phone and nunber details
name = input('What is your name? ')
speak('Hello ' + name + ' welcome to automated Python CV builder')
phone_number = input('What is your phone number? ')
email = input('What is your email? ')

document.add_paragraph(
    name + ' | ' + phone_number + " | " + email
)

# about me 
document.add_heading('About me')
document.add_paragraph(
    input('Tell me about yourself? ')
)

# work experience
document.add_heading('Work Experience')
paragraph = document.add_paragraph()

company = input('Enter company: ')
from_date = input('From date: ')
to_date = input('To date: ')

paragraph.add_run(company + ' ').bold = True
paragraph.add_run(from_date + '-' + to_date + '\n').italic = True

experinece_details = input(
    'Describe your experience at ' + company + ' '
)

paragraph.add_run(experinece_details)

# more experiences
while True:
    has_more_experiences = input(
        'Do you have more experiences?  Yes or No '
    )
    if has_more_experiences.lower() == 'yes':
        paragraph = document.add_paragraph()

        company = input('Enter company: ')
        from_date = input('From date: ')
        to_date = input('To date: ')

        paragraph.add_run(company + ' ').bold = True
        paragraph.add_run(from_date + '-' + to_date + '\n').italic = True

        experinece_details = input(
            'Describe your experience at ' + company + ' '
        )

        paragraph.add_run(experinece_details)
    else:
        break

# skills
document.add_heading('Skills')
skill = input('Enter skills: ')
paragraph = document.add_heading(skill)
paragraph.style = 'List Bullet'

while True:
    has_more_skills = input('Do you have more skills? Yes or No ')
    if has_more_skills.lower() == 'yes':
        skill = input('Enter skills: ')
        paragraph = document.add_paragraph(skill)
        paragraph.style = 'List Bullet'
    else:
        break

# footer
section = document.sections[0]
footer = section.footer
paragraph = footer.paragraphs[0]
paragraph.text = 'CV generated using automated Python CV builder'

document.save('files/cv.docx')
